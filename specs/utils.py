from docx import Document
from .models import Products
from docx.shared import Cm, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from django.core.mail import EmailMessage
import os
from django.conf import settings


def perform_data_cleaning(data):
    email_to = data['email']
    recipient_list = [email_to, ]
    name = data['customer_name']
    # model = data['model']
    print('media/attachments/'+str(data['file']))
    word_file = Document('./media/attachments/'+str(data['file']))
    # print(file_text)
    # for table in word_file.tables:

    for row in word_file.tables[0].rows:
        # for cell in row.cells:
        from_addr = row.cells[0].text
        to_addr = row.cells[1].text
    p = PlacedOrder.objects.create(shipping_address=to_addr, sending_address=from_addr)
    p.save()
    for idx, row in enumerate(word_file.tables[1].rows):
        if idx == 0:
            pass
        else:
            industry = row.cells[0].text
            model = row.cells[0].text
            size = row.cells[0].text
            price = row.cells[0].text
            product, created = Products.objects.get_or_create(industry__industry=industry, model=model, size=size, price=int(price))
            p.order_list.add(product)
    x = PlacedOrder.objects.all().count() + 1000
    p.invoice_no = x
    p.save()
    subject = 'Order placed'
    message = 'Your order has been placed successfully' + 'you invoice number is '+ x
    email_from = settings.EMAIL_HOST_USER
    email_message = EmailMessage(subject, message, email_from, recipient_list)
    email_message.send()
    print('success')


def send_email_with_attachment(email_from, email_to, ship_address, carts, city, state, pincode):
    subject = 'Invoice report'
    message = ' please refer to the attachment for invoice.'
    recipient_list = []
    for a in carts:
        print(a)
    recipient_list.append(email_to)
    newDoc = Document()
    company_name = newDoc.add_heading('Cognizant')
    company_name.alignment = WD_TABLE_ALIGNMENT.CENTER
    tableAddress = newDoc.add_table(cols=2, rows=1)
    addressRow = tableAddress.rows[0]
    addressRow.height = Inches(1)
    addressRow.cells[0].text = "From : " + email_from
    addressRow.cells[1].text = "To : " + ship_address + '\n' + city + '\n' + state + '\n' + str(pincode)
    cart_data = []
    total_price = 0
    for c in carts:
        a_cart = []
        a_cart.append(c.industry.industry)
        a_cart.append(c.model)
        a_cart.append(c.size)
        a_cart.append(c.price)
        total_price += c.price
        cart_data.append(a_cart)
    heading = newDoc.add_heading('ORDER LIST')
    heading.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = len(carts)
    table = newDoc.add_table(cols=4, rows=rows + 1)
    columns = ['COMPANY', 'MODEL_NAME', 'SIZE', 'PRICE']
    i = 0
    for cells in table.rows[0].cells:
        cells.text = columns[i]
        i += 1
    print(cart_data)
    id = 0
    for row in table.rows[1:]:
        x = 0
        for cell in row.cells:
            cell.text = str(cart_data[id][x])
            x += 1
        id += 1
    newDoc.add_paragraph('\nTotal price : ' + str(total_price))
    newDoc.save('invoice.docx')
    email_message = EmailMessage(subject, message, email_from, recipient_list)
    email_message.attach_file('invoice.docx')
    email_message.send()
    os.remove('invoice.docx')
